"""
Serviciu de integrare cu Microsoft Word via python-docx.
Generare documente .docx: facturi, rapoarte, contracte.
"""
from io import BytesIO
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from erp_app import db
from erp_app.models import Invoice, Order, Client, Employee


def _add_styled_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, value in enumerate(row_data):
            table.rows[row_idx + 1].cells[col_idx].text = str(value)

    return table


def generate_invoice_doc(invoice_id):
    """Genereaza factura in format Word (.docx)."""
    invoice = db.session.get(Invoice, invoice_id)
    if not invoice:
        return None

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Title
    title = doc.add_heading("FACTURA", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Invoice details
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(f"Nr: {invoice.number}\n").bold = True
    p.add_run(f"Data: {invoice.date.strftime('%d.%m.%Y') if invoice.date else '-'}\n")
    p.add_run(f"Scadenta: {invoice.due_date.strftime('%d.%m.%Y') if invoice.due_date else '-'}\n")

    doc.add_paragraph()

    # Client info
    if invoice.client:
        client = invoice.client
        p = doc.add_paragraph()
        p.add_run("Client: ").bold = True
        p.add_run(f"{client.name}\n")
        if client.company:
            p.add_run(f"Companie: {client.company}\n")
        if client.cui:
            p.add_run(f"CUI: {client.cui}\n")
        if client.address:
            p.add_run(f"Adresa: {client.address}")
        if client.city:
            p.add_run(f", {client.city}")

    doc.add_paragraph()

    # Items table
    headers = ["Nr.", "Descriere", "Cant.", "Pret unitar", "TVA%", "Total"]
    rows = []
    for i, item in enumerate(invoice.items, 1):
        rows.append([
            str(i),
            item.description,
            f"{item.quantity:.2f}",
            f"{item.unit_price:.2f} RON",
            f"{item.vat_rate:.0f}%",
            f"{item.total:.2f} RON",
        ])
    _add_styled_table(doc, headers, rows)

    doc.add_paragraph()

    # Totals
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(f"Subtotal: {invoice.subtotal:.2f} RON\n")
    p.add_run(f"TVA: {invoice.vat_total:.2f} RON\n")
    run = p.add_run(f"TOTAL: {invoice.total:.2f} RON\n")
    run.bold = True
    run.font.size = Pt(14)

    # Status
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    status_map = {
        "draft": "CIORNA",
        "sent": "TRIMISA",
        "paid": "PLATITA",
        "overdue": "SCADENTA DEPASITA",
        "cancelled": "ANULATA",
    }
    p.add_run(f"Status: {status_map.get(invoice.status, invoice.status)}").bold = True

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run(f"Document generat automat - ERP System - {datetime.now().strftime('%d.%m.%Y %H:%M')}").font.size = Pt(8)

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def generate_order_doc(order_id):
    """Genereaza comanda in format Word (.docx)."""
    order = db.session.get(Order, order_id)
    if not order:
        return None

    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    title = doc.add_heading("COMANDA", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(f"Nr: {order.number}\n").bold = True
    p.add_run(f"Data: {order.date.strftime('%d.%m.%Y') if order.date else '-'}\n")
    p.add_run(f"Livrare: {order.delivery_date.strftime('%d.%m.%Y') if order.delivery_date else '-'}\n")

    if order.client:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("Client: ").bold = True
        p.add_run(f"{order.client.name}")
        if order.client.company:
            p.add_run(f" ({order.client.company})")

    doc.add_paragraph()

    headers = ["Nr.", "Produs", "Cantitate", "Pret unitar", "Total"]
    rows = []
    for i, item in enumerate(order.items, 1):
        rows.append([
            str(i),
            item.product.name if item.product else "",
            f"{item.quantity:.2f}",
            f"{item.unit_price:.2f} RON",
            f"{item.total:.2f} RON",
        ])
    _add_styled_table(doc, headers, rows)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"TOTAL: {order.total:.2f} RON")
    run.bold = True
    run.font.size = Pt(14)

    if order.notes:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("Observatii: ").bold = True
        p.add_run(order.notes)

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def generate_report_doc(report_type="general"):
    """Genereaza raport general in format Word."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    title = doc.add_heading("RAPORT ERP", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Generat: {datetime.now().strftime('%d.%m.%Y %H:%M')}")

    doc.add_paragraph()

    # Statistics section
    doc.add_heading("Statistici Generale", level=1)
    stats_data = [
        ["Clienti", str(Client.query.count())],
        ["Angajati", str(Employee.query.count())],
        ["Facturi", str(Invoice.query.count())],
        ["Comenzi", str(Order.query.count())],
    ]
    _add_styled_table(doc, ["Indicator", "Valoare"], stats_data)

    doc.add_paragraph()

    # Revenue
    doc.add_heading("Situatia Financiara", level=1)
    paid = db.session.query(db.func.coalesce(db.func.sum(Invoice.total), 0)).filter_by(status="paid").scalar()
    unpaid = db.session.query(db.func.coalesce(db.func.sum(Invoice.total), 0)).filter_by(status="sent").scalar()
    finance_data = [
        ["Incasari", f"{paid:.2f} RON"],
        ["De incasat", f"{unpaid:.2f} RON"],
        ["Facturi neplatite", str(Invoice.query.filter_by(status="sent").count())],
    ]
    _add_styled_table(doc, ["Indicator", "Valoare"], finance_data)

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output
